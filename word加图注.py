from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

def insert_paragraph_after(paragraph, text=None):
    """兼容旧版本 python-docx，在指定段落后插入新段落"""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    para_obj = Paragraph(new_p, paragraph._parent)
    if text:
        run = para_obj.add_run(text)
    return para_obj

def add_caption_below_images(doc_path, output_path, chapter_number=5, start_from=1):
    """
    参数说明：
    - doc_path: 原始 Word 文件路径（.docx）
    - output_path: 输出的 Word 文件路径
    - chapter_number: 图注前缀章节号（如图5-1 中的5）
    - start_from: 从第几张图片开始加图注（从1开始编号）
    """
    doc = Document(doc_path)
    count = 1
    image_index = 1  # 当前图片是第几张
    i = 0

    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        has_picture = any('graphic' in run._element.xml for run in para.runs)

        if has_picture:
            if image_index >= start_from:
                caption_text = f"图{chapter_number}-{count}"
                caption_para = insert_paragraph_after(para, caption_text)
                caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = caption_para.runs[0]
                run.font.size = Pt(12)  # 小四 = 12pt
                run.font.name = '仿宋'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 设置中文字体

                count += 1
                i += 1  # 跳过新加的图注段落
            image_index += 1  # 不管加不加图注，图片编号都要加
        i += 1

    doc.save(output_path)
    print(f"✅ 图注添加完成，从第 {start_from} 张开始，共添加 {count - 1} 条。输出文件: {output_path}")

add_caption_below_images(
    doc_path="input.docx",
    output_path="output_with_captions.docx",
    chapter_number=5,
    start_from=12  # 从第3张图开始加图注
)